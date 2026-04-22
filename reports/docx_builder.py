# -*- coding: utf-8 -*-
"""
Builder de dossiê DOCX — gera relatório técnico profissional em Word.
Utiliza python-docx para criar tabelas formatadas, cores e seções estruturadas.
"""
from pathlib import Path
from datetime import datetime
from decimal import Decimal

from api.models.sped_file import AnalysisResult
from api.models.finding import Finding, Severity, BlockSummary


# Importações condicionais (python-docx pode não estar instalado)
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# Cores do Design System
COLOR_PRIMARY = RGBColor(0x6C, 0x5C, 0xE7) if HAS_DOCX else None  # Roxo
COLOR_EXCELLENT = RGBColor(0x04, 0x78, 0x57) if HAS_DOCX else None # Esmeralda
COLOR_SUCCESS = RGBColor(0x00, 0xB8, 0x94) if HAS_DOCX else None  # Verde
COLOR_REASONABLE = RGBColor(0xCA, 0x8A, 0x04) if HAS_DOCX else None # Âmbar
COLOR_WARNING = RGBColor(0xFD, 0xCB, 0x6E) if HAS_DOCX else None  # Amarelo
COLOR_DANGER = RGBColor(0xE1, 0x7C, 0x80) if HAS_DOCX else None   # Vermelho
COLOR_INADEQUATE = RGBColor(0x99, 0x1B, 0x1B) if HAS_DOCX else None # Carmesim
COLOR_INFO = RGBColor(0x74, 0xB9, 0xFF) if HAS_DOCX else None     # Azul
COLOR_DARK = RGBColor(0x2D, 0x3A, 0x4A) if HAS_DOCX else None     # Escuro
COLOR_LIGHT = RGBColor(0xF8, 0xF9, 0xFA) if HAS_DOCX else None    # Claro
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF) if HAS_DOCX else None

# Emoji/indicadores de severidade
SEVERITY_ICONS = {
    Severity.CRITICAL: "❌",
    Severity.WARNING: "⚠️",
    Severity.INFO: "ℹ️",
}

SEVERITY_LABELS = {
    Severity.CRITICAL: "CRÍTICO",
    Severity.WARNING: "ATENÇÃO",
    Severity.INFO: "INFORMATIVO",
}

STATUS_ICONS = {
    "ok": "✅",
    "warning": "⚠️",
    "critical": "❌",
}


class DocxBuilder:
    """Constrói o dossiê técnico em formato DOCX."""

    def __init__(self, analysis: AnalysisResult):
        self.analysis = analysis
        self.doc = Document() if HAS_DOCX else None

    def build(self):
        """Constrói todas as seções do dossiê."""
        if not HAS_DOCX:
            raise ImportError("python-docx não instalado. Execute: pip install python-docx")

        self._setup_styles()
        self._add_header()
        self._add_quadro_resumo()
        self._add_score_section()
        self._add_block_overview()
        self._add_findings_section()
        self._add_cross_validations()
        self._add_action_plan()
        self._add_footer()

    def save(self, path: Path):
        """Salva o documento DOCX."""
        self.doc.save(str(path))

    def _setup_styles(self):
        """Configura estilos do documento."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.color.rgb = COLOR_DARK

        # Reduzir margens
        for section in self.doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

    def _add_header(self):
        """Adiciona cabeçalho do dossiê."""
        # Título
        title = self.doc.add_heading('DOSSIÊ TÉCNICO — EFD Compliance', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = COLOR_PRIMARY

        # Subtítulo
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Validação Expert de SPED EFD ICMS/IPI — Análise Pós-PVA')
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_DARK
        run.font.italic = True

        self.doc.add_paragraph()  # Espaço

    def _add_quadro_resumo(self):
        """Adiciona quadro-resumo com dados do contribuinte."""
        self.doc.add_heading('1. Dados do Contribuinte', level=1)

        info = self.analysis.file_info
        table = self.doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        dados = [
            ("Razão Social", info.razao_social),
            ("CNPJ", self._format_cnpj(info.cnpj)),
            ("Inscrição Estadual", info.ie),
            ("UF", info.uf),
            ("Período", f"{info.periodo_ini.strftime('%d/%m/%Y') if info.periodo_ini else 'N/D'} a {info.periodo_fin.strftime('%d/%m/%Y') if info.periodo_fin else 'N/D'}"),
            ("Perfil", f"Perfil {info.perfil}" if info.perfil else "N/D"),
            ("Versão do Layout", info.cod_ver),
            ("Arquivo", self.analysis.filename),
        ]

        for i, (label, valor) in enumerate(dados):
            cells = table.rows[i].cells
            cells[0].text = label
            cells[1].text = str(valor) if valor else "N/D"
            # Negrito no label
            for paragraph in cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    def _add_score_section(self):
        """Adiciona seção do score de conformidade."""
        self.doc.add_heading('2. Score de Conformidade', level=1)

        score = self.analysis.score
        if score >= 91:
            label = "EXCELENTE"
            color = COLOR_EXCELLENT
        elif score >= 71:
            label = "BOM"
            color = COLOR_SUCCESS
        elif score >= 51:
            label = "RAZOÁVEL"
            color = COLOR_REASONABLE
        elif score >= 41:
            label = "ATENÇÃO"
            color = COLOR_WARNING
        elif score >= 21:
            label = "CRÍTICO"
            color = COLOR_DANGER
        else:
            label = "INADEQUADO"
            color = COLOR_INADEQUATE

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{score:.1f}% — {label}')
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = color

        # Totalizadores
        total_c = self.analysis.total_critical
        total_w = self.analysis.total_warnings
        total_i = self.analysis.total_info

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(f'❌ {total_c} Críticos   ').font.color.rgb = COLOR_DANGER
        p2.add_run(f'⚠️ {total_w} Atenção   ').font.color.rgb = COLOR_WARNING if HAS_DOCX else None
        p2.add_run(f'ℹ️ {total_i} Informativos').font.color.rgb = COLOR_INFO

    def _add_block_overview(self):
        """Adiciona visão geral por bloco."""
        self.doc.add_heading('3. Visão por Bloco', level=1)

        table = self.doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Cabeçalho
        headers = ['Bloco', 'Descrição', 'Registros', 'Achados', 'Status']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            for p in table.rows[0].cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True

        # Dados
        for code, summary in sorted(self.analysis.block_summaries.items()):
            row = table.add_row()
            row.cells[0].text = code
            row.cells[1].text = summary.block_name
            row.cells[2].text = str(summary.total_records)
            row.cells[3].text = str(summary.total_findings)
            
            p_status = row.cells[4].paragraphs[0]
            run_status = p_status.add_run(STATUS_ICONS.get(summary.status, ""))
            if summary.status == "ok":
                run_status.font.color.rgb = COLOR_SUCCESS
            elif summary.status == "warning":
                run_status.font.color.rgb = COLOR_WARNING
            elif summary.status == "critical":
                run_status.font.color.rgb = COLOR_DANGER

    def _add_findings_section(self):
        """Adiciona seção detalhada de achados."""
        self.doc.add_heading('4. Achados Detalhados', level=1)

        if not self.analysis.findings:
            self.doc.add_paragraph('Nenhum achado identificado. ✅')
            return

        # Agrupar por severidade
        for severity in [Severity.CRITICAL, Severity.WARNING, Severity.INFO]:
            findings = [f for f in self.analysis.findings if f.severity == severity]
            if not findings:
                continue

            icon = SEVERITY_ICONS[severity]
            label = SEVERITY_LABELS[severity]
            self.doc.add_heading(f'{icon} {label} ({len(findings)})', level=2)

            for finding in findings:
                # Título do achado
                p = self.doc.add_paragraph()
                run = p.add_run(f'[{finding.code}] {finding.title}')
                run.bold = True
                run.font.size = Pt(10)

                # Descrição
                if finding.description:
                    self.doc.add_paragraph(finding.description)

                # Valores esperado vs encontrado
                if finding.expected_value or finding.actual_value:
                    t = self.doc.add_table(rows=1, cols=2)
                    t.rows[0].cells[0].text = f"Esperado: {finding.expected_value or 'N/D'}"
                    t.rows[0].cells[1].text = f"Encontrado: {finding.actual_value or 'N/D'}"

                # Referência legal
                if finding.legal_reference:
                    p_ref = self.doc.add_paragraph()
                    run_ref = p_ref.add_run(f'📖 {finding.legal_reference}')
                    run_ref.font.size = Pt(8)
                    run_ref.font.italic = True
                    run_ref.font.color.rgb = COLOR_INFO

                # Recomendação
                if finding.recommendation:
                    p_rec = self.doc.add_paragraph()
                    run_rec = p_rec.add_run(f'💡 {finding.recommendation}')
                    run_rec.font.size = Pt(9)

                self.doc.add_paragraph()  # Separador

    def _add_cross_validations(self):
        """Adiciona seção de verificações cruzadas."""
        cross_findings = [f for f in self.analysis.findings if "×" in f.register]
        if not cross_findings:
            return

        self.doc.add_heading('5. Verificações Cruzadas entre Blocos', level=1)

        for finding in cross_findings:
            p = self.doc.add_paragraph()
            icon = SEVERITY_ICONS[finding.severity]
            run = p.add_run(f'{icon} [{finding.code}] {finding.title}')
            run.bold = True

            if finding.description:
                self.doc.add_paragraph(finding.description)

    def _add_action_plan(self):
        """Adiciona plano de ação com prioridades."""
        critical = [f for f in self.analysis.findings if f.severity == Severity.CRITICAL]
        warnings = [f for f in self.analysis.findings if f.severity == Severity.WARNING]

        if not critical and not warnings:
            return

        self.doc.add_heading('6. Plano de Ação', level=1)

        if critical:
            self.doc.add_heading('Ações Imediatas (Críticas)', level=2)
            for i, f in enumerate(critical, 1):
                rec = f.recommendation or "Verificar e corrigir conforme diagnóstico."
                self.doc.add_paragraph(f'{i}. [{f.code}] {rec}', style='List Number')

        if warnings:
            self.doc.add_heading('Ações Recomendadas (Atenção)', level=2)
            for i, f in enumerate(warnings, 1):
                rec = f.recommendation or "Analisar e avaliar necessidade de correção."
                self.doc.add_paragraph(f'{i}. [{f.code}] {rec}', style='List Number')

    def _add_footer(self):
        """Adiciona rodapé com metadados."""
        self.doc.add_paragraph()
        
        # Linha separadora sólida e profissional no rodapé
        p_line = self.doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_line = p_line.add_run('―' * 55)
        run_line.font.color.rgb = RGBColor(0xDE, 0xE2, 0xE6) # Cinza claro sólido

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        now = datetime.now().strftime("%d/%m/%Y %H:%M")
        run = p.add_run(
            f'Dossiê gerado por EFD Compliance v{self.analysis.validator_version}\n'
            f'Data: {now}\n'
            f'Hash do arquivo: {self.analysis.file_hash[:16]}...'
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    @staticmethod
    def _format_cnpj(cnpj: str) -> str:
        """Formata CNPJ: XX.XXX.XXX/XXXX-XX"""
        cnpj = cnpj.replace(".", "").replace("/", "").replace("-", "")
        if len(cnpj) == 14:
            return f"{cnpj[:2]}.{cnpj[2:5]}.{cnpj[5:8]}/{cnpj[8:12]}-{cnpj[12:]}"
        return cnpj
